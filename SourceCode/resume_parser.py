"""
resume_parser.py — Resume Document Parser
Extracts structured information from PDF/DOCX resumes using
regex patterns and spaCy NLP for entity recognition.
"""

import re
import io
from dataclasses import dataclass, field
from typing import Optional

# PDF parsing
import pdfplumber

# DOCX parsing
from docx import Document


@dataclass
class ParsedResume:
    """Structured representation of an extracted resume."""
    raw_text: str = ""
    name: str = ""
    email: str = ""
    phone: str = ""
    skills: list[str] = field(default_factory=list)
    experience_years: float = 0.0
    experience_text: str = ""
    education_level: str = ""  # PhD, Masters, Bachelors, Diploma, etc.
    education_text: str = ""
    job_titles: list[str] = field(default_factory=list)
    notice_period_days: int = 90  # default assumption
    certifications: list[str] = field(default_factory=list)
    summary: str = ""


# ─── Section Header Patterns ──────────────────────────────────
SECTION_PATTERNS = {
    "skills": re.compile(
        r"(?i)^(?:technical\s+)?skills|competenc|technologies|proficienc|expertise",
        re.MULTILINE
    ),
    "experience": re.compile(
        r"(?i)^(?:work\s+)?experience|employment|career\s+history|professional\s+background|work\s+history",
        re.MULTILINE
    ),
    "education": re.compile(
        r"(?i)^education|academic|qualification|degree",
        re.MULTILINE
    ),
    "summary": re.compile(
        r"(?i)^(?:professional\s+)?summary|objective|profile|about\s+me|career\s+objective",
        re.MULTILINE
    ),
    "certifications": re.compile(
        r"(?i)^certif|licens|accredit|awards",
        re.MULTILINE
    ),
}

# ─── Skill Extraction Patterns ────────────────────────────────
TECH_SKILLS_DB = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl",
    # Frameworks
    "react", "angular", "vue", "django", "flask", "fastapi", "spring",
    "node.js", "express", "next.js", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "spark", "hadoop", "airflow",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "dynamodb", "cassandra", "neo4j", "oracle",
    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins",
    "ci/cd", "git", "linux", "ansible", "grafana", "prometheus",
    # AI/ML
    "machine learning", "deep learning", "nlp", "computer vision",
    "transformers", "bert", "gpt", "llm", "reinforcement learning",
    "neural networks", "random forest", "xgboost", "svm",
    "natural language processing", "generative ai", "rag",
    # Data
    "data science", "data engineering", "etl", "data pipeline",
    "power bi", "tableau", "looker", "data visualization",
    "statistics", "a/b testing", "hypothesis testing",
    # Other
    "agile", "scrum", "jira", "rest api", "graphql", "microservices",
    "system design", "oop", "design patterns", "html", "css",
]

# ─── Education Level Mapping ──────────────────────────────────
EDUCATION_LEVELS = {
    "phd": "PhD",
    "ph.d": "PhD",
    "doctorate": "PhD",
    "doctor of": "PhD",
    "master": "Masters",
    "mtech": "Masters",
    "m.tech": "Masters",
    "msc": "Masters",
    "m.sc": "Masters",
    "mba": "Masters",
    "m.s.": "Masters",
    "bachelor": "Bachelors",
    "b.tech": "Bachelors",
    "btech": "Bachelors",
    "b.sc": "Bachelors",
    "bsc": "Bachelors",
    "b.e.": "Bachelors",
    "b.eng": "Bachelors",
    "diploma": "Diploma",
    "associate": "Diploma",
    "certificate": "Certificate",
    "high school": "HighSchool",
    "secondary": "HighSchool",
}

EDUCATION_ORDINAL = {
    "PhD": 5,
    "Masters": 4,
    "Bachelors": 3,
    "Diploma": 2,
    "Certificate": 1,
    "HighSchool": 0,
    "Unknown": 0,
}

# ─── Job Title Patterns ───────────────────────────────────────
JOB_TITLE_PATTERNS = [
    r"(?i)(?:senior|junior|lead|principal|staff|chief)?\s*(?:software|data|ml|ai|devops|cloud|full\s*stack|front\s*end|back\s*end|mobile)\s*(?:engineer|developer|architect|scientist|analyst)",
    r"(?i)(?:senior|junior|lead|principal)?\s*(?:product|project|program|engineering)\s*manager",
    r"(?i)(?:senior|junior)?\s*(?:business|data|systems?|research)\s*analyst",
    r"(?i)(?:cto|ceo|coo|vp|director|head)\s+(?:of\s+)?(?:engineering|technology|data|product)",
    r"(?i)(?:technical|team|engineering)\s*lead",
    r"(?i)(?:solutions?|enterprise|technical)\s*architect",
    r"(?i)data\s*(?:engineer|scientist|analyst)",
    r"(?i)machine\s*learning\s*engineer",
    r"(?i)research\s*(?:engineer|scientist)",
    r"(?i)consultant",
    r"(?i)intern",
]


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


def extract_sections(text: str) -> dict[str, str]:
    """Split resume text into labelled sections using header patterns."""
    lines = text.split("\n")
    sections = {}
    current_section = "header"
    current_lines = []

    for line in lines:
        matched = False
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.search(line.strip()):
                # Save previous section
                sections[current_section] = "\n".join(current_lines)
                current_section = section_name
                current_lines = []
                matched = True
                break
        if not matched:
            current_lines.append(line)

    # Save last section
    sections[current_section] = "\n".join(current_lines)
    return sections


def extract_email(text: str) -> str:
    """Extract email address from text."""
    match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract phone number from text."""
    match = re.search(r"[\+]?[(]?\d{1,4}[)]?[-\s./]?\d{3,4}[-\s./]?\d{4}", text)
    return match.group(0) if match else ""


def extract_skills(text: str) -> list[str]:
    """Match known technical skills against resume text."""
    text_lower = text.lower()
    found_skills = []
    for skill in TECH_SKILLS_DB:
        # Use word boundary matching for short skills to avoid false positives
        if len(skill) <= 3:
            if re.search(r"\b" + re.escape(skill) + r"\b", text_lower):
                found_skills.append(skill)
        else:
            if skill in text_lower:
                found_skills.append(skill)
    return list(set(found_skills))


def extract_experience_years(text: str) -> float:
    """Estimate total years of experience from text patterns."""
    # Pattern 1: "X years of experience" or "X+ years"
    matches = re.findall(
        r"(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)",
        text, re.IGNORECASE
    )
    if matches:
        return max(float(m) for m in matches)

    # Pattern 2: Count distinct year ranges (2019-2023 = 4 years)
    year_ranges = re.findall(r"(20\d{2})\s*[-–—to]+\s*(20\d{2}|present|current)",
                             text, re.IGNORECASE)
    total = 0
    for start, end in year_ranges:
        start_y = int(start)
        end_y = 2026 if end.lower() in ("present", "current") else int(end)
        total += max(0, end_y - start_y)

    return min(total, 40)  # cap at 40 years


def extract_education_level(text: str) -> str:
    """Determine highest education level mentioned."""
    text_lower = text.lower()
    highest = "Unknown"
    highest_ord = -1

    for keyword, level in EDUCATION_LEVELS.items():
        if keyword in text_lower:
            ord_val = EDUCATION_ORDINAL.get(level, 0)
            if ord_val > highest_ord:
                highest_ord = ord_val
                highest = level

    return highest


def extract_job_titles(text: str) -> list[str]:
    """Extract job titles from resume text."""
    titles = []
    for pattern in JOB_TITLE_PATTERNS:
        matches = re.findall(pattern, text)
        titles.extend(matches)
    # Clean and deduplicate
    cleaned = list(set(t.strip().title() for t in titles if len(t.strip()) > 3))
    return cleaned[:5]  # limit to top 5


def extract_notice_period(text: str) -> int:
    """Extract notice period in days from resume text."""
    text_lower = text.lower()

    # Pattern: "notice period: X days/weeks/months"
    match = re.search(
        r"notice\s*period[:\s]*(\d+)\s*(day|week|month)",
        text_lower
    )
    if match:
        value = int(match.group(1))
        unit = match.group(2)
        if "week" in unit:
            return value * 7
        elif "month" in unit:
            return value * 30
        return value

    # Pattern: "immediately available" / "available immediately"
    if re.search(r"immediat(?:ely)?\s*available|available\s*immediat", text_lower):
        return 0

    # Default assumption
    return 90


def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
    """
    Main entry point: parse a resume file into structured data.
    Supports PDF and DOCX formats.
    """
    # Step 1: Extract raw text based on file type
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raw_text = file_bytes.decode("utf-8", errors="ignore")

    if not raw_text.strip():
        return ParsedResume(raw_text="[Empty document]")

    # Step 2: Split into sections
    sections = extract_sections(raw_text)

    # Step 3: Extract structured fields
    header_text = sections.get("header", "")
    skills_text = sections.get("skills", "") + " " + raw_text
    experience_text = sections.get("experience", "")
    education_text = sections.get("education", "")

    # Extract name from first non-empty line of header
    name_candidates = [
        line.strip() for line in header_text.split("\n")
        if line.strip() and not re.search(r"@|http|www|\d{5,}", line)
    ]
    name = name_candidates[0] if name_candidates else "Unknown"
    # Clean name: remove common non-name elements
    name = re.sub(r"(?i)curriculum\s+vitae|resume|cv", "", name).strip()

    resume = ParsedResume(
        raw_text=raw_text,
        name=name[:100],
        email=extract_email(raw_text),
        phone=extract_phone(raw_text),
        skills=extract_skills(skills_text),
        experience_years=extract_experience_years(raw_text),
        experience_text=experience_text[:2000],
        education_level=extract_education_level(education_text or raw_text),
        education_text=education_text[:1000],
        job_titles=extract_job_titles(raw_text),
        notice_period_days=extract_notice_period(raw_text),
        certifications=[],
        summary=sections.get("summary", "")[:500],
    )

    return resume


if __name__ == "__main__":
    # Quick test
    sample = """
    John Doe
    john.doe@email.com | +65 9123 4567

    PROFESSIONAL SUMMARY
    Senior Data Scientist with 8 years of experience in machine learning and NLP.

    TECHNICAL SKILLS
    Python, TensorFlow, PyTorch, SQL, AWS, Docker, Kubernetes, React

    WORK EXPERIENCE
    Senior Data Scientist | Google | 2020 - Present
    - Built NLP pipelines processing 1M documents daily

    Data Scientist | Facebook | 2016 - 2020
    - Developed recommendation engine using deep learning

    EDUCATION
    Master of Science in Computer Science | Stanford University | 2016
    Bachelor of Science in Mathematics | MIT | 2014

    Notice Period: 30 days
    """
    result = parse_resume(sample.encode(), "test.txt")
    print(f"Name: {result.name}")
    print(f"Skills: {result.skills}")
    print(f"Experience: {result.experience_years} years")
    print(f"Education: {result.education_level}")
    print(f"Job Titles: {result.job_titles}")
    print(f"Notice Period: {result.notice_period_days} days")
