"""
resume_parser.py — Resume Document Parser
Extracts structured information from PDF/DOCX resumes using
regex patterns and spaCy NLP for entity recognition.
"""

import re
import io
import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional

# PDF parsing
import pdfplumber

# DOCX parsing
from docx import Document


SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}
DEFAULT_RESUME_DATASET_DIR = (
    Path(__file__).resolve().parents[3] / "dataset" / "test_dataset" / "AGRICULTURE"
)


@dataclass
class ParsedResume:
    """Structured representation of an extracted resume."""
    raw_text: str = ""
    name: str = ""
    email: str = ""
    phone: str = ""
    skills: list[str] = field(default_factory=list)
    experience_years: float = 0.0
    experience_text: str = ""
    education_level: str = ""  # PhD, Masters, Bachelors, Diploma, etc.
    education_text: str = ""
    job_titles: list[str] = field(default_factory=list)
    notice_period_days: int = 90  # default assumption
    certifications: list[str] = field(default_factory=list)
    career_gaps: list[dict] = field(default_factory=list)  # gaps between roles >= 12 months
    summary: str = ""


# ─── Section Header Patterns ──────────────────────────────────
SECTION_PATTERNS = {
    "skills": re.compile(
        r"(?i)^(?:(?:technical\s*)?skills|competenc(?:y|ies)?|technologies|proficienc(?:y|ies)?|expertise)$",
        re.MULTILINE
    ),
    "experience": re.compile(
        r"(?i)^(?:(?:work\s*)?experience|employment|career\s*history|professional\s*background|work\s*history)$",
        re.MULTILINE
    ),
    "education": re.compile(
        r"(?i)^(?:education|academic(?:\s*background)?|qualifications?)$",
        re.MULTILINE
    ),
    "summary": re.compile(
        r"(?i)^(?:(?:professional\s*)?summary|objective|profile|about\s*me|career\s*objective)$",
        re.MULTILINE
    ),
    "certifications": re.compile(
        r"(?i)^(?:certif\w*|licens\w*|accredit\w*|awards?)$",
        re.MULTILINE
    ),
}

EDUCATION_BLOCK_LINE_PATTERNS = [
    re.compile(r"(?i)\b(?:bachelor(?:'s)?|master(?:'s)?|doctorate|ph\.?d\.?|mba|mtech|m\.tech|msc|m\.sc|btech|b\.tech|bsc|b\.sc|beng|b\.eng|associate\s+degree|associate\s+of|diploma)\b"),
    re.compile(r"(?i)\b(?:university|college|polytechnic|institute|school\s+of|faculty\s+of)\b"),
]

EDUCATION_BLOCK_REJECT_PATTERNS = [
    re.compile(r"(?i)\b(?:certified|certification|certificate|foundations?|badge|credential|scrum|oracle|microsoft|python\s+institute|azure|power\s+bi)\b"),
]

EDUCATION_BLOCK_STOP_PATTERNS = [
    SECTION_PATTERNS["skills"],
    SECTION_PATTERNS["experience"],
    SECTION_PATTERNS["summary"],
    SECTION_PATTERNS["certifications"],
]

EXPERIENCE_DATE_PATTERN = re.compile(
    r"(?i)(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[\s\-]*\d{4}\s*[-–—]\s*(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[\s\-]*\d{4}|present|ongoing|current|now)|\b20\d{2}\s*[-–—]\s*(?:20\d{2}|present|ongoing|current)\b"
)

EXPERIENCE_STOP_PATTERNS = [
    SECTION_PATTERNS["education"],
    SECTION_PATTERNS["certifications"],
    re.compile(r"(?i)^awards?(?:\s*and\s*achievements?)?$") ,
    re.compile(r"(?i)^projects?$"),
    re.compile(r"(?i)^technical\s*skills$") ,
]

ASSOCIATE_ACADEMIC_CONTEXT = [
    "associate degree",
    "associate of",
    "of science",
    "of arts",
    "community college",
    "college",
    "university",
    "polytechnic",
    "institute",
    "school",
]

# ─── Skill Extraction Patterns ────────────────────────────────
TECH_SKILLS_DB = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl",
    # Frameworks
    "react", "angular", "vue", "django", "flask", "fastapi", "spring",
    "node.js", "express", "next.js", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "spark", "hadoop", "airflow",
    "testng", "junit", "bdd", "cucumber", "selenium",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "dynamodb", "cassandra", "neo4j", "oracle", "plsql", "pl/sql",
    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins",
    "ci/cd", "git", "linux", "ansible", "grafana", "prometheus",
    "jboss", "gradle", "maven", "tomcat",
    # AI/ML
    "machine learning", "deep learning", "nlp", "computer vision",
    "transformers", "bert", "gpt", "llm", "reinforcement learning",
    "neural networks", "random forest", "xgboost", "svm",
    "natural language processing", "generative ai", "rag",
    # Data
    "data science", "data engineering", "etl", "data pipeline",
    "power bi", "tableau", "looker", "data visualization",
    "statistics", "a/b testing", "hypothesis testing",
    # General SDLC / methodology
    "agile", "waterfall", "scrum", "jira", "rest api", "graphql", "microservices",
    "system design", "oop", "oops", "design patterns", "html", "css",
    "kanban", "scrum master", "communication", "problem solving",
    "open api", "open banking",
    # Domain — banking / finance core
    "banking", "finance", "fintech", "compliance", "audit", "risk management",
    "investment banking", "capital markets", "asset management",
    # Domain — trading systems / Calypso ecosystem
    "calypso", "calypso architecture", "calypso modules", "calypso data model",
    "trading platform", "trading platforms", "trade lifecycle",
    "front office", "back office", "middle office",
    "settlements", "trade settlement", "clearing",
    "p&l", "official p&l", "accounting setup",
    "market data", "curve setup", "static data",
    "position management", "position configuration",
    "bloomberg sapi", "bloomberg", "reuters", "toms", "deal tracker",
    "market data interface",
    # Domain — risk / derivatives
    "ers", "limit setup", "var", "value at risk",
    "credit risk", "market risk", "operational risk",
    "fixed income", "money market", "interest rate", "interest rate derivatives",
    "fx", "fx derivatives", "foreign exchange",
    "credit derivatives", "non linear derivatives", "non-linear derivatives",
    "derivatives", "swaps", "options", "futures", "bonds",
    "asset class", "swift messaging",
    # Other industries
    "healthcare", "e-commerce", "cybersecurity", "blockchain",
    # Soft skills (frequently appear in JDs)
    "leadership", "stakeholder management", "test management", "uat",
    "automation", "manual testing",
]

# ─── Education Level Mapping ──────────────────────────────────
EDUCATION_LEVELS = {
    "phd": "PhD",
    "ph.d": "PhD",
    "doctorate": "PhD",
    "doctor of": "PhD",
    "masters of": "Masters",
    "mtech": "Masters",
    "m.tech": "Masters",
    "msc": "Masters",
    "m.sc": "Masters",
    "mba": "Masters",
    "m.s.": "Masters",
    "bachelor": "Bachelors",
    "b.tech": "Bachelors",
    "btech": "Bachelors",
    "b.sc": "Bachelors",
    "bsc": "Bachelors",
    "b.e.": "Bachelors",
    "b.eng": "Bachelors",
    "diploma": "Diploma",
    "associate": "Diploma",
    "certificate": "Certificate",
    "high school": "HighSchool",
    "HSC": "HighSchool",
    "secondary": "HighSchool",
    "SSC": "HighSchool",
}

EDUCATION_ORDINAL = {
    "PhD": 5,
    "Masters": 4,
    "Bachelors": 3,
    "Diploma": 2,
    "HighSchool": 1,
    "Unknown": 0,
}

# ─── Job Title Patterns ───────────────────────────────────────
JOB_TITLE_PATTERNS = [
    # ─── Tech / Engineering ───
    r"(?i)(?:senior|junior|lead|principal|staff|chief)?\s*(?:software|data|ml|ai|devops|cloud|full\s*stack|front\s*end|back\s*end|mobile|systems?|security|network)\s*(?:engineer|developer|architect|scientist|analyst|administrator)",
    r"(?i)(?:senior|junior|lead|principal)?\s*(?:product|project|program|engineering)\s*manager",
    r"(?i)(?:senior|junior)?\s*(?:business|data|systems?|research|financial|qa|test)\s*analyst",
    r"(?i)(?:cto|ceo|coo|cfo|vp|director|head)\s+(?:of\s+)?(?:engineering|technology|data|product|operations|sales)",
    r"(?i)(?:technical|team|engineering|tech)\s*lead",
    r"(?i)(?:solutions?|enterprise|technical|cloud|security|data)\s*architect",
    r"(?i)data\s*(?:engineer|scientist|analyst)",
    r"(?i)machine\s*learning\s*engineer",
    r"(?i)research\s*(?:engineer|scientist)",
    r"(?i)\b(?:consultant|intern|trainee|associate|specialist)\b",
    # ─── QA / Testing ───
    r"(?i)(?:senior|junior|lead|principal)?\s*(?:test|qa|quality\s*assurance)\s*(?:lead|manager|engineer|analyst|architect|automation)",
    r"(?i)(?:senior|junior|lead)?\s*(?:automation|sdet|performance)\s*(?:engineer|tester|analyst)",
    r"(?i)test\s*architect",
    r"(?i)uat\s*(?:test\s*)?(?:manager|lead|analyst)",
    # ─── Generic management ───
    r"(?i)(?:senior|junior|lead|head)?\s*(?:program|project|product|account|operations)\s*manager",
    # ─── Non-tech professions (CRITICAL for eligibility detection) ───
    r"(?i)(?:yoga|fitness|pilates|zumba|aerobics|reformer)\s*(?:teacher|instructor|trainer|coach)",
    r"(?i)(?:personal|gym|sports|athletic)\s*trainer",
    r"(?i)(?:executive|head|sous|pastry|line|station)\s*chef",
    r"(?i)\b(?:chef|cook|baker|sommelier)\b",
    r"(?i)\b(?:cashier|waiter|waitress|bartender|barista|hostess?|server)\b",
    r"(?i)(?:hairdresser|hair\s*stylist|barber|beautician|cosmetologist|makeup\s*artist|esthetician|nail\s*technician)",
    r"(?i)(?:plumber|electrician|carpenter|welder|mason|painter|roofer|locksmith|machinist)",
    r"(?i)(?:gardener|florist|landscaper|farmer|farmhand)",
    r"(?i)(?:truck|cab|taxi|delivery|bus|chauffeur)\s*driver",
    r"(?i)(?:security\s*guard|bouncer|doorman|watchman)",
    r"(?i)\b(?:nurse|nursing|caregiver|midwife|paramedic|medic|doctor|surgeon|dentist|veterinarian|pharmacist)\b",
    r"(?i)(?:physical|occupational|speech)\s*therapist",
    r"(?i)group\s*fitness\s*(?:instructor|trainer)",
    r"(?i)(?:retail|store|shop)\s*(?:associate|clerk|assistant)",
    r"(?i)(?:teacher|tutor|professor|lecturer|educator|instructor)",  # education roles
    r"(?i)(?:lawyer|attorney|paralegal|legal\s*assistant|advocate|barrister|solicitor)",
    r"(?i)(?:journalist|reporter|editor|copywriter|content\s*writer)",
]


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


def extract_sections(text: str) -> dict[str, str]:
    """Split resume text into labelled sections using header patterns."""
    lines = text.split("\n")
    sections = {}
    current_section = "header"
    current_lines = []

    for line in lines:
        matched = False
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.search(line.strip()):
                # Save previous section
                sections[current_section] = "\n".join(current_lines)
                current_section = section_name
                current_lines = []
                matched = True
                break
        if not matched:
            current_lines.append(line)

    # Save last section
    sections[current_section] = "\n".join(current_lines)
    return sections


def extract_education_block(text: str) -> str:
    """Recover education lines when the resume has no explicit education header."""
    lines = [line.strip() for line in text.splitlines()]
    collected_lines = []
    collecting = False

    for line in lines:
        if not line:
            if collecting and collected_lines:
                break
            continue

        if any(pattern.search(line) for pattern in EDUCATION_BLOCK_STOP_PATTERNS):
            if collecting:
                break
            continue

        line_matches_education = any(pattern.search(line) for pattern in EDUCATION_BLOCK_LINE_PATTERNS)
        if line_matches_education and any(pattern.search(line) for pattern in EDUCATION_BLOCK_REJECT_PATTERNS):
            line_matches_education = False
        if line_matches_education:
            collecting = True
            collected_lines.append(line)
            continue

        if collecting:
            # Keep nearby continuation lines such as institution, year, or location.
            if len(line) <= 120:
                collected_lines.append(line)
                continue
            break

    return "\n".join(collected_lines)


def extract_experience_block(text: str) -> str:
    """Recover experience lines when the resume has no explicit experience header."""
    lines = [line.strip() for line in text.splitlines()]
    collected_lines = []
    collecting = False

    for index, line in enumerate(lines):
        if not line:
            if collecting and collected_lines:
                break
            continue

        if any(pattern.search(line) for pattern in EXPERIENCE_STOP_PATTERNS):
            if collecting:
                break
            continue

        if not collecting and EXPERIENCE_DATE_PATTERN.search(line):
            collecting = True
            previous_line = lines[index - 1].strip() if index > 0 else ""
            if previous_line and len(previous_line) <= 80 and not EXPERIENCE_DATE_PATTERN.search(previous_line):
                collected_lines.append(previous_line)
            collected_lines.append(line)
            continue

        if collecting:
            collected_lines.append(line)

    return "\n".join(collected_lines)


def extract_email(text: str) -> str:
    """Extract email address from text."""
    match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract phone number from text."""
    match = re.search(r"[\+]?[(]?\d{1,4}[)]?[-\s./]?\d{3,4}[-\s./]?\d{4}", text)
    return match.group(0) if match else ""


def clean_output_text(text: str) -> str:
    """Normalize output text to alphanumeric characters and whitespace only."""
    cleaned_lines = []
    for line in (text or "").splitlines():
        cleaned = re.sub(r"[^A-Za-z0-9\s]", " ", line)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if cleaned:
            cleaned_lines.append(cleaned)
    return "\n".join(cleaned_lines)


def extract_skills(text: str) -> list[str]:
    """Match known skills against resume text.

    Uses three strategies (in priority order):
      1. Literal word-boundary match of skill in text.
      2. Substring match for multi-word skills.
      3. Domain-inference rules — captures abstract/soft skills that rarely
         appear as literal terms in resumes. Examples:
           - 'banking' inferred from 'Bank of America', 'BNP Paribas', 'JPMorgan'
           - 'communication' inferred from 'communicating effectively with stakeholders'
           - 'problem solving' inferred from 'analytical', 'troubleshooting'

    The third strategy is critical because resumes rarely contain literal
    terms like 'banking' or 'communication' — those are expressed indirectly.
    """
    text_lower = text.lower()
    found_skills = set()

    # Strategy 1 + 2: literal and word-boundary matches
    for skill in TECH_SKILLS_DB:
        skill_l = skill.lower()
        if len(skill_l) <= 3:
            if re.search(r"\b" + re.escape(skill_l) + r"\b", text_lower):
                found_skills.add(skill_l)
        else:
            if skill_l in text_lower:
                found_skills.add(skill_l)

    # Strategy 3: domain-inference — credit a skill if its concept is
    # strongly implied even when the literal word isn't present.
    SKILL_INFERENCE_RULES = {
        "banking": [
            r"(?i)\b(?:bank\s+of\s+\w+|banking|jpmorgan|hsbc|citibank|"
            r"standard\s+chartered|deutsche\s+bank|barclays|bnp\s+paribas|"
            r"credit\s+agricole|abn\s+amro|goldman\s+sachs|morgan\s+stanley|"
            r"wells\s+fargo|state\s+street|swift\s+messag|trade\s+process|"
            r"forex|foreign\s+exchange|securities|bond\s+asset|bfsi|"
            r"capital\s+market|investment\s+bank)\b"
        ],
        "communication": [
            r"(?i)\b(?:communicat|stakeholder|presentation|reporting|liais|"
            r"correspond|interpersonal|verbal\s+and\s+written|cross[-\s]functional)",
        ],
        "problem solving": [
            r"(?i)\b(?:problem[-\s]solv|analytical|troubleshoot|root[-\s]cause|"
            r"debugg|critical[-\s]think|investigat|diagnos)",
        ],
        "leadership": [
            r"(?i)\b(?:leader|managed|mentored|coached|led\s+(?:a|the|team)|"
            r"team\s+lead|head\s+of|directed)",
        ],
        "agile": [
            r"(?i)\b(?:agile|scrum|sprint|kanban|standup|retrospective|"
            r"product\s+owner|backlog\s+groom)",
        ],
        "test management": [
            r"(?i)\b(?:test\s+(?:management|strategy|plan|estimate|architect)|"
            r"qa\s+lead|test\s+lead|defect\s+manage|sign[-\s]off)",
        ],
        "uat": [
            r"(?i)\buat\b|user\s+acceptance\s+test",
        ],
        "automation": [
            r"(?i)\b(?:automat|cucumber|webdriver|rest[-\s]assured|playwright)",
        ],
        "finance": [
            r"(?i)\b(?:finance|financial|treasury|accounting|audit|fund\s+management|"
            r"asset\s+management|portfolio|hedge\s+fund)",
        ],
        "stakeholder management": [
            r"(?i)\b(?:stakeholder|client\s+management|business\s+user|liaison)",
        ],
    }
    for skill, patterns in SKILL_INFERENCE_RULES.items():
        if skill in found_skills:
            continue
        for pat in patterns:
            if re.search(pat, text):
                found_skills.add(skill)
                break

    return list(found_skills)


def _is_education_context(text: str, match_start: int, match_end: int) -> bool:
    """Check if a date range is sitting within a degree/diploma line.

    Strategy: look at the line containing the match. If that line (or the
    line immediately before it) contains a degree keyword, the date range
    is most likely a degree timeline rather than work experience.

    Word boundaries are essential — 'mba' must NOT match inside 'mumbai',
    'bsc' must NOT match inside other words, etc.
    """
    EDU_LINE_KEYWORDS = [
        # Degree names (the strongest signal — usually on the same line as the dates)
        "bachelor", "bachelors", "master", "masters", "phd", "doctorate", "doctoral",
        "diploma", "mba", "btech", "mtech", "bsc", "msc",
        "b\\.tech", "m\\.tech", "b\\.sc", "m\\.sc", "b\\.e\\.", "m\\.e\\.",
        # Institution words
        "university", "college", "academy",
        # Degree program markers
        "graduated", "thesis", "dissertation",
    ]
    # Phrases (handled separately — these are multi-word so word boundary
    # at the spaces in the middle would be wrong)
    EDU_LINE_PHRASES = ["institute of", "school of"]

    # Build a single word-boundary regex
    edu_pattern = r"\b(?:" + "|".join(EDU_LINE_KEYWORDS) + r")\b"

    # Find the line containing this match
    line_start = text.rfind("\n", 0, match_start) + 1
    line_end = text.find("\n", match_end)
    if line_end == -1:
        line_end = len(text)
    line = text[line_start:line_end].lower()

    # Check the same line first
    if re.search(edu_pattern, line, re.IGNORECASE):
        return True
    if any(p in line for p in EDU_LINE_PHRASES):
        return True

    # Also check the line immediately before — sometimes 'EDUCATION' header
    # sits on its own line and the degree+date is on the next line.
    if line_start > 0:
        prev_line_end = line_start - 1  # the newline char
        prev_line_start = text.rfind("\n", 0, prev_line_end) + 1
        prev_line = text[prev_line_start:prev_line_end].lower()
        if re.search(edu_pattern, prev_line, re.IGNORECASE):
            return True
        if any(p in prev_line for p in EDU_LINE_PHRASES):
            return True
        # Also check for the bare EDUCATION header
        if "education" in prev_line and len(prev_line.strip()) <= 30:
            return True

    return False


def extract_experience_years(text: str, experience_section: Optional[str] = None) -> float:
    """Estimate total years of experience from text patterns.

    Strategies in priority order:
      1. Explicit statement like "11 years of experience" (most reliable).
         Handles spacing variants: "11years", "11 yrs", "11+ years", etc.
      2. Month + year ranges within the experience section: "Aug 2023 - Ongoing",
         "Jun-2018 - Jan-2020". De-duplicates overlapping ranges.
      3. Plain year ranges as last-resort fallback (e.g. "2019 - 2023") —
         applied ONLY to the experience section so we don't count degree dates
         like "Bachelor of Engineering 2010 - 2014".

    Two layers of education-date protection:
      - The caller passes only the experience section in `experience_section`
        when it can.
      - Even within that section, date ranges in an education-context window
        (degree keywords nearby) are skipped — handles cases where section
        detection failed.

    Args:
        text: Full resume text.
        experience_section: If provided, Pattern 2 and 3 use this instead of `text`
            so we don't sum work years with education years.
    """
    # ─── Pattern 1: explicit "X years of experience" ──────
    # More permissive — handle "11years" (no space), "11+ years", "11 yrs of exp"
    pattern1 = (
        r"(\d{1,2})\+?\s*(?:years?|yrs?)"          # number + unit
        r"\s*(?:of\s+)?"                             # optional "of"
        r"(?:experience|exp\b|industry|professional|relevant)"  # context word
    )
    matches = re.findall(pattern1, text, re.IGNORECASE)
    if matches:
        # Take the MAX explicit claim — handles "10+ years experience in ML, 5 years in NLP"
        years = [float(m) for m in matches if 0 < float(m) <= 50]
        if years:
            return min(40.0, max(years))

    # Also try "X+ years" alone if it's right at the top of the resume (header summary)
    header = text[:500]
    summary_match = re.search(r"(\d{1,2})\+\s*(?:years?|yrs?)\b", header, re.IGNORECASE)
    if summary_match:
        years = float(summary_match.group(1))
        if 0 < years <= 50:
            return min(40.0, years)

    # ─── Restrict patterns 2 and 3 to the experience section ──
    # If no experience section was passed, use full text (less accurate but workable).
    search_text = experience_section if experience_section else text

    # ─── Pattern 2: month + year ranges ───────────────────
    MONTH = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)"
    range_re = re.compile(
        rf"{MONTH}[\s\-]*?(\d{{4}})\s*[-–—]\s*"
        rf"(?:{MONTH}[\s\-]*?(\d{{4}})|(Present|Ongoing|Current|Now))",
        re.IGNORECASE,
    )
    ranges = []
    for m in range_re.finditer(search_text):
        # Skip if this date range is in an education-context window
        if _is_education_context(search_text, m.start(), m.end()):
            continue
        start_year = int(m.group(1))
        if m.group(2):
            end_year = int(m.group(2))
        else:
            end_year = 2026  # Present/Ongoing
        if 1990 <= start_year <= 2030 and end_year >= start_year:
            ranges.append((start_year, end_year))

    if ranges:
        # De-overlap: merge overlapping/contiguous ranges before summing
        ranges.sort()
        merged = [ranges[0]]
        for s, e in ranges[1:]:
            ls, le = merged[-1]
            if s <= le:  # overlap or contiguous
                merged[-1] = (ls, max(le, e))
            else:
                merged.append((s, e))
        total = sum(e - s for s, e in merged)
        if total > 0:
            return min(40.0, float(total))

    # ─── Pattern 3: plain "YYYY - YYYY" (last resort) ─────
    plain_re = re.compile(
        r"(20\d{2})\s*[-–—]\s*(20\d{2}|present|current|ongoing)",
        re.IGNORECASE,
    )
    pr = []
    for m in plain_re.finditer(search_text):
        # Skip if this date range is in an education-context window
        if _is_education_context(search_text, m.start(), m.end()):
            continue
        start_y = int(m.group(1))
        end_str = m.group(2)
        end_y = 2026 if end_str.lower() in ("present", "current", "ongoing") else int(end_str)
        if start_y < end_y:
            pr.append((start_y, end_y))

    if pr:
        # De-overlap
        pr.sort()
        merged = [pr[0]]
        for s, e in pr[1:]:
            ls, le = merged[-1]
            if s <= le:
                merged[-1] = (ls, max(le, e))
            else:
                merged.append((s, e))
        total = sum(e - s for s, e in merged)
        return min(40.0, float(total))

    return 0.0


def extract_education_level(text: str) -> str:
    """Determine highest education level mentioned.

    Uses word-boundary regex to avoid:
      - 'scrum master' matching 'master'
      - 'project associate' matching 'associate'
      - 'phpdev' matching 'phd'
      - 'masterclass' matching 'master'

    Also rejects matches whose immediate context is clearly NOT
    an academic degree (e.g. 'scrum master', 'master of ceremonies',
    'master class', 'master degree program' is OK).
    """
    text_lower = text.lower()
    highest = "Unknown"
    highest_ord = -1

    # Words that, when adjacent to "master"/"associate"/etc., indicate
    # a NON-academic context. We reject the match if these are nearby.
    REJECT_NEAR_MASTER = [
        "scrum", "project", "product", "story", "ceremony", "ceremonies",
        "class", "key", "mind", "yoga", "chess", "game",
    ]
    REJECT_NEAR_ASSOCIATE = [
        "project", "junior", "senior", "sales", "marketing", "research",
        "executive", "team",
    ]
    REJECT_NEAR_PHD = [
        # rare false positives, mostly safe
    ]

    # Build patterns with word boundaries
    # Order matters: longer/more-specific keys first
    education_keys = [
        ("phd", "PhD"),
        ("ph.d", "PhD"),
        ("ph\\.d\\.", "PhD"),
        ("doctorate", "PhD"),
        ("doctor of philosophy", "PhD"),
        ("master of", "Masters"),
        ("master's", "Masters"),
        ("masters degree", "Masters"),
        ("master degree", "Masters"),
        ("m\\.tech", "Masters"),
        ("mtech", "Masters"),
        ("m\\.sc", "Masters"),
        ("msc", "Masters"),
        ("m\\.s\\.", "Masters"),
        ("m\\.eng", "Masters"),
        ("meng", "Masters"),
        ("mba", "Masters"),
        ("master", "Masters"),  # fallback — applies REJECT list
        ("bachelor of", "Bachelors"),
        ("bachelor's", "Bachelors"),
        ("bachelors degree", "Bachelors"),
        ("bachelor degree", "Bachelors"),
        ("b\\.tech", "Bachelors"),
        ("btech", "Bachelors"),
        ("b\\.sc", "Bachelors"),
        ("bsc", "Bachelors"),
        ("b\\.e\\.", "Bachelors"),
        ("b\\.eng", "Bachelors"),
        ("beng", "Bachelors"),
        ("bachelor", "Bachelors"),
        ("diploma", "Diploma"),
        ("associate of", "Diploma"),
        ("associate degree", "Diploma"),
        ("associate", "Diploma"),  # applies REJECT list
        ("certificate", "Certificate"),
        ("high school", "HighSchool"),
        ("secondary school", "HighSchool"),
    ]

    for keyword_pat, level in education_keys:
        # word-boundary regex
        pat = r"\b" + keyword_pat + r"\b"
        for m in re.finditer(pat, text_lower):
            # Pull a small window around the match to inspect context
            start = max(0, m.start() - 25)
            end = min(len(text_lower), m.end() + 25)
            context = text_lower[start:end]

            # Apply rejection rules
            matched_word = m.group(0)
            if matched_word == "master":
                if any(reject in context for reject in REJECT_NEAR_MASTER):
                    continue
            if matched_word == "associate":
                if any(reject in context for reject in REJECT_NEAR_ASSOCIATE):
                    continue
                if not any(marker in context for marker in ASSOCIATE_ACADEMIC_CONTEXT):
                    continue

            # Accept this match
            ord_val = EDUCATION_ORDINAL.get(level, 0)
            if ord_val > highest_ord:
                highest_ord = ord_val
                highest = level
            break  # one match per keyword is enough

    return highest


def extract_job_titles(text: str) -> list[str]:
    """Extract job titles from resume text."""
    titles = []
    for pattern in JOB_TITLE_PATTERNS:
        matches = re.findall(pattern, text)
        titles.extend(matches)
    # Clean and deduplicate
    cleaned = list(set(t.strip().title() for t in titles if len(t.strip()) > 3))
    return cleaned[:5]  # limit to top 5


def extract_notice_period(text: str) -> int:
    """Extract notice period in days from resume text."""
    text_lower = text.lower()

    # Pattern: "notice period: X days/weeks/months"
    match = re.search(
        r"notice\s*period[:\s]*(\d+)\s*(day|week|month)",
        text_lower
    )
    if match:
        value = int(match.group(1))
        unit = match.group(2)
        if "week" in unit:
            return value * 7
        elif "month" in unit:
            return value * 30
        return value

    # Pattern: "immediately available" / "available immediately"
    if re.search(r"immediat(?:ely)?\s*available|available\s*immediat", text_lower):
        return 0

    # Default assumption
    return 90


def extract_certifications(text: str, certifications_section: str = "") -> list[str]:
    """Extract professional certifications from the resume.

    Looks for well-known certification patterns (industry-standard names) AND
    free-form lines from a 'Certifications' section. Returns a deduplicated
    list of certification names found.

    Search priority:
      1. The dedicated CERTIFICATIONS section if present (each line is a cert).
      2. Body-text matching against a curated list of well-known certs.
      3. Generic 'X Certified' / 'Certified X' patterns.

    The Misc scorer uses len(certifications) as a reward signal.
    """
    found = set()

    # ─── Pass 1: well-known cert patterns anywhere in the resume ──
    # These are the most common professional certifications we see.
    KNOWN_CERTS = [
        # Cloud
        ("AWS Certified", r"(?i)\baws\s+certified(?:\s+\w+){0,4}"),
        ("AWS Solutions Architect", r"(?i)\baws\s+(?:solutions?\s+)?architect"),
        ("Azure Certified", r"(?i)\bazure\s+(?:certified|fundamentals|administrator|architect|engineer)"),
        ("GCP Certified", r"(?i)\bgcp\s+(?:certified|professional|associate)|google\s+cloud\s+(?:certified|professional)"),
        # Project / Scrum / Agile
        ("PMP", r"(?i)\bpmp\b|project\s+management\s+professional"),
        ("PRINCE2", r"(?i)\bprince2\b"),
        ("CSM", r"(?i)\bcsm\b|certified\s+scrum\s+master"),
        ("CSPO", r"(?i)\bcspo\b|certified\s+scrum\s+product\s+owner"),
        ("SAFe", r"(?i)\bsafe(?:\s+agilist|\s+practitioner)?\b"),
        # Testing / QA
        ("ISTQB", r"(?i)\bistqb\b"),
        ("CSTE", r"(?i)\bcste\b"),
        # Banking / Finance
        ("CFA", r"(?i)\bcfa\b"),
        ("FRM", r"(?i)\bfrm\b"),
        ("CPA", r"(?i)\bcpa\b"),
        ("Calypso Certified", r"(?i)\bcalypso\s+(?:certified|certification)"),
        # IT / Security
        ("CISSP", r"(?i)\bcissp\b"),
        ("CISA", r"(?i)\bcisa\b"),
        ("CompTIA", r"(?i)\bcomptia\s+(?:a\+|network\+|security\+)"),
        ("CCNA", r"(?i)\bccna\b"),
        ("CCNP", r"(?i)\bccnp\b"),
        # Data
        ("Tableau Certified", r"(?i)\btableau\s+(?:certified|specialist|associate)"),
        ("Power BI Certified", r"(?i)\bpower\s+bi\s+(?:certified|data\s+analyst)"),
        ("Databricks Certified", r"(?i)\bdatabricks\s+certified"),
        # Java / Oracle / Microsoft
        ("OCJP", r"(?i)\bocjp\b|oracle\s+certified\s+java"),
        ("Microsoft Certified", r"(?i)\bmicrosoft\s+certified"),
        # Awards/recognition treated as certification-like
        ("Six Sigma", r"(?i)\bsix\s+sigma\b|\b(?:green|black)\s+belt\b"),
        ("ITIL", r"(?i)\bitil\b"),
    ]

    for cert_name, pattern in KNOWN_CERTS:
        if re.search(pattern, text):
            found.add(cert_name)

    # ─── Pass 2: parse the dedicated certifications section line-by-line ──
    if certifications_section:
        for line in certifications_section.splitlines():
            line = line.strip()
            # Skip empty lines and section headers
            if not line or len(line) < 4 or line.lower().startswith(("certif", "licens")):
                continue
            # Skip dates and bullets
            cleaned = re.sub(r"^[\-•·*▪◦◆●■▲★]\s*", "", line)
            cleaned = re.sub(r"\s*\([\d/\-]+\)\s*$", "", cleaned)  # strip "(2023)"
            cleaned = re.sub(r"\s*[-–]\s*\d{4}\s*$", "", cleaned)  # strip "- 2023"
            if 4 < len(cleaned) < 120:
                # Cap to first ~80 chars to avoid grabbing entire paragraphs
                found.add(cleaned[:80].strip())

    # ─── Pass 3: generic "X Certified" / "Certified X" fallback ──
    # Catches certifications not in the curated list. Anchored to a single line
    # so we don't capture newlines into the cert name.
    for m in re.finditer(r"(?i)\b(?:certified|certification)[\s:,-]+([A-Z][^\n\r]{3,40})", text):
        candidate = m.group(1).strip()
        # Remove trailing punctuation / continuation markers
        candidate = re.sub(r"[.,;:].*$", "", candidate).strip()
        # Avoid noise like "Certified by ..." or generic words
        if candidate and not candidate.lower().startswith(("by ", "in ", "from ", "and ")):
            found.add(candidate[:50])

    # Limit to 15 to avoid runaway counts
    return sorted(found)[:15]


def detect_career_gaps(text: str, experience_section: str = "",
                        threshold_months: int = 12) -> list[dict]:
    """Detect gaps between consecutive work experience entries.

    Walks date ranges in chronological order. Any gap larger than `threshold_months`
    months (default 12 = 1 year) between the END of one role and the START of the
    next is flagged.

    Returns a list of dicts: {"from_year": int, "to_year": int, "gap_months": int}.
    The Misc scorer penalises candidates with one or more such gaps.

    Note: this only catches gaps BETWEEN roles. Time before the first role or
    after the last role is not counted as a gap (those are pre-career and
    current-unemployment respectively).
    """
    search_text = experience_section or text
    if not search_text:
        return []

    # Reuse the existing month+year range regex
    MONTH = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)"
    range_re = re.compile(
        rf"({MONTH})[\s\-]*?(\d{{4}})\s*[-–—]\s*"
        rf"(?:({MONTH})[\s\-]*?(\d{{4}})|(Present|Ongoing|Current|Now))",
        re.IGNORECASE,
    )

    MONTH_NUM = {
        "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
        "jul": 7, "aug": 8, "sep": 9, "sept": 9, "oct": 10, "nov": 11, "dec": 12,
    }

    ranges = []  # list of (start_month_index, end_month_index)
    for m in range_re.finditer(search_text):
        # Skip education-context dates
        if _is_education_context(search_text, m.start(), m.end()):
            continue

        start_mo = MONTH_NUM.get(m.group(1).lower(), 1)
        start_yr = int(m.group(2))

        if m.group(3):
            end_mo = MONTH_NUM.get(m.group(3).lower(), 12)
            end_yr = int(m.group(4))
        else:
            # Present/Ongoing
            end_mo = 12
            end_yr = 2026

        if not (1990 <= start_yr <= 2030):
            continue

        start_idx = start_yr * 12 + start_mo
        end_idx = end_yr * 12 + end_mo
        if end_idx >= start_idx:
            ranges.append((start_idx, end_idx))

    if len(ranges) < 2:
        return []

    # Sort by start, then walk and find gaps
    ranges.sort()
    gaps = []
    prev_end = ranges[0][1]
    for start, end in ranges[1:]:
        gap = start - prev_end
        if gap >= threshold_months:
            gaps.append({
                "from_year": prev_end // 12,
                "to_year": start // 12,
                "gap_months": gap,
            })
        prev_end = max(prev_end, end)

    return gaps


def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
    """
    Main entry point: parse a resume file into structured data.
    Supports PDF and DOCX formats.
    """
    # Step 1: Extract raw text based on file type
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raw_text = file_bytes.decode("utf-8", errors="ignore")

    if not raw_text.strip():
        return ParsedResume(raw_text="[Empty document]")

    # Step 2: Split into sections
    sections = extract_sections(raw_text)

    # Step 3: Extract structured fields
    header_text = sections.get("header", "")
    skills_text = sections.get("skills", "") + " " + raw_text
    experience_text = sections.get("experience", "")
    if len(experience_text.strip()) <= 10:
        experience_text = extract_experience_block(raw_text)
    education_text = sections.get("education", "")
    if len(education_text.strip()) <= 10:
        education_text = extract_education_block(raw_text)

    # Extract name from first non-empty line of header
    name_candidates = [
        line.strip() for line in header_text.split("\n")
        if line.strip() and not re.search(r"@|http|www|\d{5,}", line)
    ]
    name = name_candidates[0] if name_candidates else "Unknown"
    # Clean name: remove common non-name elements
    name = re.sub(r"(?i)curriculum\s+vitae|resume|cv", "", name).strip()

    # Education: STRONGLY prefer the EDUCATION section.
    # This avoids matching "Scrum Master Certified" in skills as "Masters degree".
    # Only fall back to the full text if no EDUCATION section was detected.
    if education_text and len(education_text.strip()) > 10:
        education_level = extract_education_level(education_text)
        # If the dedicated section yielded nothing, scan summary too (some CVs
        # mention "Master's degree" only in the summary, not under EDUCATION)
        if education_level == "Unknown":
            summary_text = sections.get("summary", "")
            education_level = extract_education_level(education_text + " " + summary_text)
    else:
        # No education section — fall back to full text but with the strict matcher
        education_level = extract_education_level(raw_text)

    if education_level == "Unknown":
        education_level = "Bachelors"

    # Build the experience-search text. Strip education content out of it so
    # date ranges like "Bachelor of Engineering Jun 2010 – Jul 2014" never
    # get counted as work experience. We pass this stripped text as the
    # experience_section argument to extract_experience_years.
    if experience_text and education_text:
        # Use the dedicated experience section, minus any education text
        # that may have leaked into it (defensive double-strip).
        clean_experience_text = experience_text.replace(education_text, "")
    elif experience_text:
        clean_experience_text = experience_text
    elif education_text:
        # No experience section detected — fall back to raw_text minus education
        clean_experience_text = raw_text.replace(education_text, "")
    else:
        clean_experience_text = raw_text

    # Extract certifications from the dedicated section (if present) plus
    # body-text scanning. The Misc scorer rewards each cert.
    certifications_section = sections.get("certifications", "")
    certifications = extract_certifications(raw_text, certifications_section)

    # Detect career gaps — gaps of 12+ months between roles will trigger
    # a Misc penalty. Walks date ranges in the cleaned experience text.
    career_gaps = detect_career_gaps(raw_text, clean_experience_text)

    resume = ParsedResume(
        raw_text=raw_text,
        name=clean_output_text(name)[:100] or "Unknown",
        email=extract_email(raw_text),
        phone=extract_phone(raw_text),
        skills=extract_skills(skills_text),
        experience_years=extract_experience_years(raw_text, clean_experience_text),
        experience_text=clean_output_text(experience_text)[:2000],
        education_level=education_level,
        education_text=clean_output_text(education_text)[:1000],
        job_titles=[clean_output_text(title) for title in extract_job_titles(raw_text)],
        notice_period_days=extract_notice_period(raw_text),
        certifications=certifications,
        career_gaps=career_gaps,
        summary=clean_output_text(sections.get("summary", ""))[:500],
    )

    return resume


def parse_resumes_from_directory(directory: Optional[str | Path] = None) -> list[ParsedResume]:
    """Parse only resume files from the configured dataset directory."""
    resume_dir = Path(directory) if directory else DEFAULT_RESUME_DATASET_DIR
    if not resume_dir.exists():
        raise FileNotFoundError(f"Resume directory not found: {resume_dir}")

    parsed_resumes = []
    for file_path in sorted(resume_dir.iterdir()):
        if file_path.suffix.lower() not in SUPPORTED_RESUME_EXTENSIONS:
            continue

        parsed_resumes.append(parse_resume(file_path.read_bytes(), file_path.name))

    return parsed_resumes


if __name__ == "__main__":
    results = parse_resumes_from_directory()
    print(f"Parsed {len(results)} resumes from {DEFAULT_RESUME_DATASET_DIR}")
    for resume in results:
        print(json.dumps(asdict(resume), indent=2))
